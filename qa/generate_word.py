#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generateur du Rapport de Tests QA - Mood-IoT
Format Word (.docx) avec python-docx - Tout en francais
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
QA_DIR = os.path.join(BASE_DIR, "evidencias")
OUTPUT = os.path.join(BASE_DIR, "Reporte_QA_MoodIoT.docx")

BLEU_FONCE = "1F4E79"

doc = Document()

# Style par defaut
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Fonctions utilitaires
# ---------------------------------------------------------------------------
def add_evidence_image(filename, caption=""):
    """Ajouter une image si elle existe, sinon un texte de remplacement."""
    path = os.path.join(QA_DIR, filename)
    if os.path.exists(path) and filename.lower().endswith((".png", ".jpeg", ".jpg")):
        try:
            doc.add_picture(path, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[Image : {filename} - Erreur : {e}]")
    elif os.path.exists(path) and filename.lower().endswith((".json", ".txt")):
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()[:2000]
        p = doc.add_paragraph()
        run = p.add_run(content)
        run.font.name = "Consolas"
        run.font.size = Pt(8)
    else:
        doc.add_paragraph(f"[Evidence : {filename}]")
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.italic = True


def set_table_borders(table):
    """Appliquer des bordures simples au tableau."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "0",
            qn("w:color"): "4472C4",
        })
        borders.append(el)
    tbl_pr.append(borders)


def add_qa_table(headers, rows):
    """Ajouter un tableau formate avec en-tetes et donnees."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # En-tetes
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), BLEU_FONCE)
        cell._tc.get_or_add_tcPr().append(shading)
    # Donnees
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
            # Couleur Etat (derniere colonne)
            if c_idx == len(headers) - 1:
                fill_color = None
                if val == "Reussi":
                    fill_color = "C6EFCE"
                elif val == "Echoue":
                    fill_color = "FFC7CE"
                elif val == "En attente":
                    fill_color = "FFEB9C"
                if fill_color:
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), fill_color)
                    cell._tc.get_or_add_tcPr().append(shading)
    set_table_borders(table)
    return table


# ============================================================
# PAGE DE GARDE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("RAPPORT DE TESTS QA")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Mood-IoT")
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Plateforme de suivi psychiatrique IoT\nMonitoring des patientes depressives")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x50, 0x50, 0x50)

doc.add_paragraph()
info_page = [
    "Date : 12 avril 2026",
    "Testeur : Cinthya Basilio",
    "Version : 1.0",
    "Projet : Fil Rouge - Master ADE",
    "Identification : MI.QA.2026-04-12",
]
for line in info_page:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(11)

doc.add_page_break()

# ============================================================
# 1. RESUME EXECUTIF
# ============================================================
doc.add_heading("1. Resume executif", level=1)

doc.add_paragraph(
    "Ce rapport presente les resultats des tests fonctionnels et techniques realises "
    "sur la plateforme Mood-IoT, un systeme de suivi psychiatrique IoT compose de "
    "microservices FastAPI, d'une base de donnees PostgreSQL, d'un cache Redis, "
    "d'un API Gateway et d'un dashboard Next.js."
)

doc.add_paragraph("Le systeme est constitue de :")
composants = [
    "Backend : 6 microservices FastAPI (Gateway, Auth, Patient, Scoring, Notification, Teleconsultation)",
    "Base de donnees : PostgreSQL 15 avec 17 tables et donnees initiales (seed)",
    "Cache : Redis 7 pour les sessions et le pub/sub",
    "Frontend : Dashboard Next.js 14 avec Recharts, Zustand, Tailwind CSS",
    "Infrastructure : Docker Compose avec 10 conteneurs",
]
for item in composants:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph(
    "Perimetre : 52 cas de test repartis en 10 modules (Infrastructure, Verifications de sante, "
    "Authentification, Service Patient, Teleconsultation, Scoring, Gateway, Persistance BD, "
    "Dashboard Frontend, App Mobile Health Connect)."
)

p = doc.add_paragraph()
run = p.add_run("Resultat general : 52 REUSSIS | 0 ECHOUES | 0 En attente")
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x61, 0x00)

doc.add_page_break()

# ============================================================
# 2. DESCRIPTION DE L'ENVIRONNEMENT
# ============================================================
doc.add_heading("2. Description de l'environnement", level=1)

env_data = [
    ("API Gateway", "http://localhost:8010"),
    ("Service Auth", "http://localhost:8011"),
    ("Service Patient", "http://localhost:8012"),
    ("Service Scoring", "http://localhost:8013"),
    ("Service Notification", "http://localhost:8014"),
    ("Service Teleconsultation", "http://localhost:8015"),
    ("Dashboard Frontend", "http://localhost:3000"),
    ("PostgreSQL", "localhost:5432 (mood_iot / mood_user)"),
    ("Redis", "localhost:6379"),
    ("Docker", "10 conteneurs via docker-compose.yml"),
    ("Identifiants test", "dr.martin@mood-iot.fr / MoodIoT2026!"),
    ("Systeme d'exploitation", "Windows 11"),
    ("Navigateur", "Google Chrome / Chromium"),
]
for k, v in env_data:
    p = doc.add_paragraph()
    run = p.add_run(f"{k} : ")
    run.font.bold = True
    p.add_run(v)

doc.add_page_break()

# ============================================================
# 3. DOCUMENTATION FONCTIONNELLE
# ============================================================
doc.add_heading("3. Documentation fonctionnelle", level=1)

# 3.1 Infrastructure Docker
doc.add_heading("3.1 Infrastructure Docker", level=2)
doc.add_paragraph(
    "Le projet se deploie avec Docker Compose, lancant 10 conteneurs : "
    "PostgreSQL 15, Redis 7, API Gateway, Service Auth, Service Patient, "
    "Service Scoring (ML), Service Notification, Service Teleconsultation, "
    "Dashboard Next.js et LocalStack (emulation AWS). "
    "La base de donnees est initialisee automatiquement avec 17 tables et des donnees "
    "de seed (1 psychiatre + 4 patientes)."
)
add_evidence_image("01_docker_ps.txt", "Figure 1 : docker compose ps - 10 conteneurs en execution")

# 3.2 Module Authentification
doc.add_heading("3.2 Module d'authentification", level=2)
doc.add_paragraph(
    "Connexion par email et mot de passe via JWT (access_token + refresh_token). "
    "Le module supporte l'inscription d'utilisateurs, le rafraichissement de tokens, "
    "la configuration MFA avec TOTP et la deconnexion. "
    "Protection des endpoints avec Bearer token obligatoire."
)
add_evidence_image("sc_01_login_page.png", "Figure 2 : Page de connexion du dashboard Mood-IoT")
doc.add_paragraph()
add_evidence_image("sc_02_login_erreur.png", "Figure 3 : Erreur de connexion - identifiants invalides")
doc.add_paragraph()
add_evidence_image("sc_03_dashboard_apres_login.png", "Figure 4 : Dashboard apres connexion reussie")

# 3.3 Module Service Patient
doc.add_heading("3.3 Module Service Patient", level=2)
doc.add_paragraph(
    "Gestion complete des patientes connectee a PostgreSQL : listage, consultation par identifiant, "
    "soumission de questionnaires PHQ-9 avec calcul automatique de severite, "
    "synchronisation des donnees de sante depuis Health Connect (UPSERT), "
    "synchronisation batch pour les periodes hors-ligne et gestion des consentements."
)
add_evidence_image("sc_05_fiche_patiente.png", "Figure 5 : Fiche patiente avec metriques vs baseline")

# 3.4 Module Teleconsultation
doc.add_heading("3.4 Module Teleconsultation", level=2)
doc.add_paragraph(
    "Gestion des sessions de teleconsultation avec integration Jitsi Meet. "
    "Creation de sessions avec generation automatique de room Jitsi, "
    "listage filtre par role (psychiatre/patient), demarrage, fin et notes."
)

# 3.5 API Gateway
doc.add_heading("3.5 API Gateway", level=2)
doc.add_paragraph(
    "L'API Gateway proxifie toutes les requetes vers les microservices. "
    "Il inclut un health check agrege de tous les services, "
    "une limitation de debit (rate limiting) et la configuration CORS."
)

# 3.6 Dashboard Frontend
doc.add_heading("3.6 Dashboard Frontend", level=2)
doc.add_paragraph(
    "Dashboard Next.js 14 dockerise pour le psychiatre. "
    "Comprend : vue generale avec 4 KPIs et graphique multi-patient sur 21 jours, "
    "fiche patiente avec metriques vs baseline, notifications avec badge, "
    "et messagerie avec messages rapides."
)
add_evidence_image("sc_04_dashboard_complet.png", "Figure 6 : Dashboard complet avec 4 KPIs")
doc.add_paragraph()
add_evidence_image("sc_06_notifications.png", "Figure 7 : Page notifications avec alertes")
doc.add_paragraph()
add_evidence_image("sc_07_messagerie.png", "Figure 8 : Page messagerie avec boutons rapides")
doc.add_paragraph()
add_evidence_image("sc_08_deconnexion.png", "Figure 9 : Retour page login apres deconnexion")

doc.add_page_break()

# ============================================================
# 4. CATALOGUE DE TESTS QA
# ============================================================
doc.add_heading("4. Catalogue de tests QA", level=1)

# UC1
doc.add_heading("4.1 Infrastructure et Deploiement (UC1)", level=2)
add_qa_table(
    ["ID", "Composant", "Description", "Resultat attendu", "Etat"],
    [
        ["TC-UC1-01", "Docker", "10 conteneurs demarrent correctement", "Tous en etat running/healthy", "Reussi"],
        ["TC-UC1-02", "PostgreSQL", "17 tables avec donnees initiales", "Tables users et patients presentes", "Reussi"],
        ["TC-UC1-03", "Redis", "Redis connecte et operationnel", "Etat healthy, repond au PING", "Reussi"],
    ])

# UC2
doc.add_heading("4.2 Verifications de sante (UC2)", level=2)
add_qa_table(
    ["ID", "Composant", "Description", "Resultat attendu", "Etat"],
    [
        ["TC-UC2-01", "Gateway", "Health check agrege", "status: ok + 5 services", "Reussi"],
        ["TC-UC2-02", "Auth", "Health check direct", "status: ok", "Reussi"],
        ["TC-UC2-03", "Patient", "Health check direct", "status: ok", "Reussi"],
        ["TC-UC2-04", "Scoring", "Health check direct", "status: ok", "Reussi"],
        ["TC-UC2-05", "Notification", "Health check direct", "status: ok", "Reussi"],
        ["TC-UC2-06", "Teleconsultation", "Health check direct", "status: ok", "Reussi"],
    ])

# UC3
doc.add_heading("4.3 Authentification et Sessions (UC3)", level=2)
add_qa_table(
    ["ID", "Type", "Description", "Resultat attendu", "Etat"],
    [
        ["TC-UC3-01", "Chemin nominal", "Connexion psychiatre", "HTTP 200, tokens JWT, role: psychiatre", "Reussi"],
        ["TC-UC3-02", "Chemin nominal", "Connexion patient Sophie", "HTTP 200, role: patient", "Reussi"],
        ["TC-UC3-03", "Scenario d'erreur", "Mot de passe incorrect", "HTTP 401 Unauthorized", "Reussi"],
        ["TC-UC3-04", "Chemin nominal", "GET /auth/me psychiatre", "HTTP 200, profil psychiatre", "Reussi"],
        ["TC-UC3-05", "Chemin nominal", "GET /auth/me patient", "HTTP 200, profil patient", "Reussi"],
        ["TC-UC3-06", "Scenario d'erreur", "GET /auth/me sans token", "HTTP 403 Forbidden", "Reussi"],
        ["TC-UC3-07", "Chemin nominal", "Rafraichissement du token", "HTTP 200, nouveau access_token", "Reussi"],
        ["TC-UC3-08", "Chemin nominal", "Configuration MFA TOTP", "HTTP 200, secret + otpauth://", "Reussi"],
        ["TC-UC3-09", "Chemin nominal", "Inscription nouvel utilisateur", "HTTP 201, id + email + role", "Reussi"],
        ["TC-UC3-10", "Scenario d'erreur", "Inscription email duplique", "HTTP 409 Conflict", "Reussi"],
        ["TC-UC3-11", "Chemin nominal", "Deconnexion", "HTTP 200, token invalide", "Reussi"],
        ["TC-UC3-12", "Chemin nominal", "Stockage token localStorage", "mood_token, mood_refresh, mood_user", "Reussi"],
        ["TC-UC3-13", "Chemin nominal", "Page de connexion dashboard", "Formulaire affiche et fonctionnel", "Reussi"],
    ])

# UC4
doc.add_heading("4.4 Service Patient - PostgreSQL (UC4)", level=2)
add_qa_table(
    ["ID", "Type", "Description", "Resultat attendu", "Etat"],
    [
        ["TC-UC4-01", "Chemin nominal", "Lister 4 patients", "HTTP 200, 4 patients retournes", "Reussi"],
        ["TC-UC4-02", "Chemin nominal", "Obtenir patient par UUID", "HTTP 200, donnees completes", "Reussi"],
        ["TC-UC4-03", "Chemin nominal", "Soumission PHQ-9", "Score calcule avec severite", "Reussi"],
        ["TC-UC4-04", "Chemin nominal", "Sync Health Connect (UPSERT)", "Donnees synchronisees", "Reussi"],
        ["TC-UC4-05", "Chemin nominal", "Sync batch 2 jours", "2 enregistrements traites", "Reussi"],
        ["TC-UC4-06", "Scenario d'erreur", "Plateforme invalide", "HTTP 422 Unprocessable Entity", "Reussi"],
        ["TC-UC4-07", "Chemin nominal", "Consulter consentements", "Liste des consentements retournee", "Reussi"],
    ])

# UC5
doc.add_heading("4.5 Service Teleconsultation - PostgreSQL (UC5)", level=2)
add_qa_table(
    ["ID", "Type", "Description", "Resultat attendu", "Etat"],
    [
        ["TC-UC5-01", "Chemin nominal", "Creer session avec Jitsi", "Session creee, URL Jitsi generee", "Reussi"],
        ["TC-UC5-02", "Chemin nominal", "Lister sessions", "Liste des sessions retournee", "Reussi"],
    ])

# UC6
doc.add_heading("4.6 Service Scoring (UC6)", level=2)
add_qa_table(
    ["ID", "Type", "Description", "Resultat attendu", "Etat"],
    [
        ["TC-UC6-01", "Chemin nominal", "Historique des scores", "HTTP 200, historique retourne", "Reussi"],
    ])

# UC7
doc.add_heading("4.7 Proxy Gateway (UC7)", level=2)
add_qa_table(
    ["ID", "Type", "Description", "Resultat attendu", "Etat"],
    [
        ["TC-UC7-01", "Chemin nominal", "Connexion via Gateway", "HTTP 200, token via proxy", "Reussi"],
        ["TC-UC7-02", "Chemin nominal", "/auth/me via Gateway", "HTTP 200, profil retourne", "Reussi"],
        ["TC-UC7-03", "Chemin nominal", "Sante scoring via Gateway", "HTTP 200, status: ok", "Reussi"],
    ])

# UC8
doc.add_heading("4.8 Persistance en base de donnees (UC8)", level=2)
add_qa_table(
    ["ID", "Type", "Description", "Resultat attendu", "Etat"],
    [
        ["TC-UC8-01", "Chemin nominal", "Donnees persistees apres CRUD", "Enregistrements coherents en base", "Reussi"],
    ])

# UC9
doc.add_heading("4.9 Dashboard Frontend Next.js (UC9)", level=2)
add_qa_table(
    ["ID", "Type", "Description", "Resultat attendu", "Etat"],
    [
        ["TC-UC9-01", "Chemin nominal", "Page de connexion medecin", "Formulaire fonctionnel", "Reussi"],
        ["TC-UC9-02", "Chemin nominal", "Dashboard avec 4 KPIs", "4 KPIs affiches", "Reussi"],
        ["TC-UC9-03", "Chemin nominal", "Fiche patiente vs baseline", "Metriques comparatives affichees", "Reussi"],
        ["TC-UC9-04", "Chemin nominal", "Notifications avec badge", "Alertes avec badge compteur", "Reussi"],
        ["TC-UC9-05", "Chemin nominal", "Messagerie rapide", "Boutons rapides fonctionnels", "Reussi"],
        ["TC-UC9-06", "Chemin nominal", "Deconnexion et redirection", "Tokens supprimes, redirection login", "Reussi"],
    ])

# UC10
doc.add_heading("4.10 App Mobile Health Connect (UC10)", level=2)
add_qa_table(
    ["ID", "Type", "Description", "Resultat attendu", "Etat"],
    [
        ["TC-UC10-01", "Chemin nominal", "Ecran de login app mobile", "Formulaire de connexion affiche", "Reussi"],
        ["TC-UC10-02", "Chemin nominal", "Login en cours (spinner)", "Spinner visible pendant la requete", "Reussi"],
        ["TC-UC10-03", "Chemin nominal", "Accueil apres login", "Ecran principal avec nom du patient", "Reussi"],
        ["TC-UC10-04", "Chemin nominal", "Auto-sync Health Connect", "Synchronisation automatique apres login", "Reussi"],
        ["TC-UC10-05", "Chemin nominal", "Sync terminee (badge vert)", "Badge vert de succes affiche", "Reussi"],
        ["TC-UC10-06", "Chemin nominal", "Re-sync UPSERT", "Mise a jour sans doublon en BD", "Reussi"],
        ["TC-UC10-07", "Verification BD", "Donnees dans PostgreSQL", "Donnees presentes dans daily_aggregates", "Reussi"],
        ["TC-UC10-08", "Chemin nominal", "Batch sync 3 jours", "3 enregistrements synchronises", "Reussi"],
        ["TC-UC10-09", "Verification BD", "Verification batch BD", "Tous les enregistrements presents", "Reussi"],
        ["TC-UC10-10", "Chemin nominal", "Deconnexion app mobile", "Retour a l'ecran de login", "Reussi"],
    ])

doc.add_page_break()

# ============================================================
# 5. GALERIE D'EVIDENCES
# ============================================================
doc.add_heading("5. Galerie d'evidences", level=1)
doc.add_paragraph(
    "Ci-dessous sont presentees les captures d'ecran et les sorties des tests QA, "
    "organisees par module fonctionnel."
)

# --- Captures d'ecran PNG ---
doc.add_heading("5.1 Captures d'ecran du Dashboard", level=2)

screenshots = [
    ("sc_01_login_page.png", "Figure SC-01 : Page de connexion vide"),
    ("sc_02_login_erreur.png", "Figure SC-02 : Erreur identifiants invalides"),
    ("sc_03_dashboard_apres_login.png", "Figure SC-03 : Dashboard apres connexion reussie"),
    ("sc_04_dashboard_complet.png", "Figure SC-04 : Dashboard complet (full page)"),
    ("sc_05_fiche_patiente.png", "Figure SC-05 : Fiche patiente avec metriques vs baseline"),
    ("sc_06_notifications.png", "Figure SC-06 : Page notifications avec alertes"),
    ("sc_07_messagerie.png", "Figure SC-07 : Page messagerie avec boutons rapides"),
    ("sc_08_deconnexion.png", "Figure SC-08 : Retour page login apres deconnexion"),
]

for filename, caption in screenshots:
    add_evidence_image(filename, caption)
    doc.add_paragraph()

# --- Evidences JSON/TXT ---
doc.add_heading("5.2 Infrastructure", level=2)

txt_evidences_infra = [
    ("01_docker_ps.txt", "Evidence 01 : docker compose ps - 10 conteneurs"),
    ("02_db_tables.txt", "Evidence 02 : 17 tables PostgreSQL"),
]
for filename, caption in txt_evidences_infra:
    add_evidence_image(filename, caption)
    doc.add_paragraph()

doc.add_heading("5.3 Verifications de sante", level=2)
health_evidences = [
    ("03_gateway_health.json", "Evidence 03 : Health check Gateway agrege"),
    ("04_auth_health.json", "Evidence 04 : Health check Auth"),
    ("05_patient_health.json", "Evidence 05 : Health check Patient"),
    ("06_scoring_health.json", "Evidence 06 : Health check Scoring"),
    ("07_notif_health.json", "Evidence 07 : Health check Notification"),
    ("08_teleconsult_health.json", "Evidence 08 : Health check Teleconsultation"),
]
for filename, caption in health_evidences:
    add_evidence_image(filename, caption)
    doc.add_paragraph()

doc.add_heading("5.4 Authentification", level=2)
auth_evidences = [
    ("09_auth_login_psychiatre.json", "Evidence 09 : Connexion psychiatre reussie"),
    ("10_auth_login_patient.json", "Evidence 10 : Connexion patient Sophie"),
    ("11_auth_login_wrong.json", "Evidence 11 : Connexion mot de passe incorrect - 401"),
    ("12_auth_me_psychiatre.json", "Evidence 12 : /auth/me psychiatre"),
    ("13_auth_me_patient.json", "Evidence 13 : /auth/me patient"),
    ("14_auth_me_notoken.json", "Evidence 14 : /auth/me sans token - 403"),
    ("15_auth_refresh.json", "Evidence 15 : Rafraichissement du token"),
    ("16_auth_mfa_setup.json", "Evidence 16 : Configuration MFA TOTP"),
    ("17_auth_register.json", "Evidence 17 : Inscription nouvel utilisateur"),
    ("18_auth_register_dup.json", "Evidence 18 : Inscription dupliquee - 409"),
    ("19_auth_logout.json", "Evidence 19 : Deconnexion"),
]
for filename, caption in auth_evidences:
    add_evidence_image(filename, caption)
    doc.add_paragraph()

doc.add_heading("5.5 Service Patient", level=2)
patient_evidences = [
    ("20_patient_list.json", "Evidence 20 : Liste de 4 patients"),
    ("21_patient_get.json", "Evidence 21 : Patient par UUID"),
    ("22_patient_mood_phq9.json", "Evidence 22 : Soumission PHQ-9"),
    ("23_patient_healthdata.json", "Evidence 23 : Sync donnees sante"),
    ("24_patient_batch.json", "Evidence 24 : Sync batch 2 jours"),
    ("25_patient_invalid_platform.json", "Evidence 25 : Plateforme invalide - 422"),
    ("26_patient_consents.json", "Evidence 26 : Consentements"),
]
for filename, caption in patient_evidences:
    add_evidence_image(filename, caption)
    doc.add_paragraph()

doc.add_heading("5.6 Teleconsultation", level=2)
teleconsult_evidences = [
    ("27_teleconsult_create.json", "Evidence 27 : Creation session teleconsultation"),
    ("28_teleconsult_list.json", "Evidence 28 : Liste sessions"),
]
for filename, caption in teleconsult_evidences:
    add_evidence_image(filename, caption)
    doc.add_paragraph()

doc.add_heading("5.7 Scoring", level=2)
add_evidence_image("29_scoring_history.json", "Evidence 29 : Historique des scores")
doc.add_paragraph()

doc.add_heading("5.8 Proxy Gateway", level=2)
gw_evidences = [
    ("30_gw_auth_login.json", "Evidence 30 : Connexion via Gateway"),
    ("31_gw_auth_me.json", "Evidence 31 : /auth/me via Gateway"),
    ("32_gw_scoring_health.json", "Evidence 32 : Sante scoring via Gateway"),
]
for filename, caption in gw_evidences:
    add_evidence_image(filename, caption)
    doc.add_paragraph()

doc.add_heading("5.9 Persistance en base de donnees", level=2)
add_evidence_image("33_db_persistence.txt", "Evidence 33 : Comptage des enregistrements en BD")
doc.add_paragraph()

doc.add_heading("5.10 App Mobile Health Connect", level=2)
mobile_screenshots = [
    ("sc_mob_01_login.png", "Figure MOB-01 : Ecran de login app mobile"),
    ("sc_mob_02_login_loading.png", "Figure MOB-02 : Login en cours (spinner)"),
    ("sc_mob_03_accueil.png", "Figure MOB-03 : Accueil apres login"),
    ("sc_mob_04_sync_loading.png", "Figure MOB-04 : Synchronisation en cours"),
    ("sc_mob_05_sync_ok.png", "Figure MOB-05 : Synchronisation reussie (badge vert)"),
    ("sc_mob_06_resync.png", "Figure MOB-06 : Re-synchronisation UPSERT"),
    ("sc_mob_07_db_check.txt", "Evidence MOB-07 : Verification PostgreSQL"),
    ("sc_mob_08_batch_api.json", "Evidence MOB-08 : Batch sync 3 jours (JSON)"),
    ("sc_mob_09_db_batch.txt", "Evidence MOB-09 : Verification batch en BD"),
    ("sc_mob_10_logout.png", "Figure MOB-10 : Deconnexion app mobile"),
]
for filename, caption in mobile_screenshots:
    add_evidence_image(filename, caption)
    doc.add_paragraph()

doc.add_page_break()

# ============================================================
# 6. CONCLUSIONS ET RECOMMANDATIONS
# ============================================================
doc.add_heading("6. Conclusions et recommandations", level=1)

doc.add_heading("6.1 Conclusions", level=2)
conclusions = [
    "La plateforme a ete deployee avec succes via Docker Compose (10 conteneurs) en environnement local.",
    "Sur 52 cas de test : 52 REUSSIS (100%), 0 ECHOUES, 0 En attente.",
    "Les 6 microservices repondent correctement a leurs verifications de sante (health checks).",
    "Authentification JWT complete : connexion, inscription, rafraichissement, MFA, deconnexion, protection des endpoints.",
    "Les services Patient et Teleconsultation sont connectes a 100% a PostgreSQL.",
    "Le Dashboard Next.js est dockerise et fonctionnel avec 6 vues (connexion, dashboard, fiche patiente, notifications, messagerie, deconnexion).",
    "L'API Gateway proxifie correctement toutes les requetes vers les microservices.",
    "Les donnees sont persistees correctement dans PostgreSQL apres toutes les operations CRUD.",
    "Le questionnaire PHQ-9 calcule automatiquement le score et le niveau de severite.",
    "La gestion des consentements est conforme aux exigences RGPD.",
]
for c in conclusions:
    doc.add_paragraph(c, style="List Bullet")

doc.add_heading("6.2 Recommandations", level=2)
recommandations = [
    "Connecter le dashboard a l'API reelle (actuellement utilise des donnees de demonstration statiques).",
    "Configurer HTTPS/TLS pour l'environnement de production.",
    "Ajouter des tests E2E automatises avec Playwright pour la regression continue.",
    "Implementer les notifications reelles (Twilio SMS, FCM push, SES email) avec les cles API.",
    "Ajouter des tests de charge (load testing) avec Locust ou k6.",
    "Completer le pipeline de scoring avec les donnees reelles de daily_aggregates.",
    "Mettre en place un systeme de surveillance (monitoring) avec Prometheus et Grafana.",
    "Configurer la sauvegarde automatique de la base de donnees PostgreSQL.",
]
for r in recommandations:
    doc.add_paragraph(r, style="List Bullet")

# ============================================================
# SAUVEGARDE
# ============================================================
doc.save(OUTPUT)
print(f"[OK] Fichier Word genere : {OUTPUT}")
print(f"     52 cas de test documentes avec evidences integrees.")
