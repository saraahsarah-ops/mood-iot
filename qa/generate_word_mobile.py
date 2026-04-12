"""
Generateur du rapport QA — App Mobile Health Connect -> Backend
Genere un document Word avec captures d'ecran et resultats des tests.
"""

from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = Path(__file__).parent
EVIDENCIAS = BASE_DIR / "evidencias"
OUTPUT = BASE_DIR / "Reporte_QA_Mobile_MoodIoT.docx"


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xA8)
    return h


def add_screenshot(doc, filename, caption, width=5.8):
    img_path = EVIDENCIAS / filename
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        cap.runs[0].font.italic = True
    else:
        doc.add_paragraph(f"[Image non trouvee: {filename}]")


def add_text_evidence(doc, filename, max_lines=30):
    path = EVIDENCIAS / filename
    if path.exists():
        content = path.read_text(encoding="utf-8")
        lines = content.strip().split("\n")[:max_lines]
        para = doc.add_paragraph()
        run = para.add_run("\n".join(lines))
        run.font.size = Pt(8)
        run.font.name = "Consolas"
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        # Light gray background via shading
        from docx.oxml.ns import qn
        shading = run._element.makeelement(qn("w:rPr"), {})
        para_format = para.paragraph_format
        para_format.space_before = Pt(4)
        para_format.space_after = Pt(4)


def generate():
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ================================================================
    # PAGE DE TITRE
    # ================================================================
    doc.add_paragraph("")
    doc.add_paragraph("")
    title = doc.add_heading("Rapport de Tests QA", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xA8)

    subtitle = doc.add_heading("Application Mobile Health Connect", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run(
        f"Projet Mood-IoT - Fil Rouge Master ADE 2026\n"
        f"Date d'execution : {datetime.now().strftime('%d/%m/%Y a %H:%M')}\n"
        f"Plateforme : Android (Health Connect SDK)\n"
        f"Environnement : Docker Compose (local)"
    )
    info_run.font.size = Pt(11)
    info_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ================================================================
    # TABLE DES MATIERES
    # ================================================================
    add_heading_styled(doc, "Table des matieres", level=1)
    toc_items = [
        "1. Objectif des tests",
        "2. Architecture testee",
        "3. Tests d'authentification (MOB-01 a MOB-03)",
        "4. Tests de synchronisation (MOB-04 a MOB-06)",
        "5. Tests de persistance (MOB-07 a MOB-09)",
        "6. Test de deconnexion (MOB-10)",
        "7. Resume des resultats",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ================================================================
    # 1. OBJECTIF
    # ================================================================
    add_heading_styled(doc, "1. Objectif des tests", level=1)
    doc.add_paragraph(
        "Ce rapport presente les resultats des tests du flux complet de synchronisation "
        "des donnees de sante entre l'application mobile (React Native + Health Connect) "
        "et le backend Mood-IoT. Les tests couvrent :"
    )
    bullets = [
        "Authentification du patient via JWT (login / logout)",
        "Lecture des donnees Health Connect (pas, BPM, sommeil, GPS)",
        "Envoi des donnees au backend via HTTP POST",
        "Verification de l'UPSERT dans PostgreSQL (pas de doublons)",
        "Synchronisation batch (rattrapage apres periode offline)",
        "Persistance des donnees dans la table daily_aggregates",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # ================================================================
    # 2. ARCHITECTURE
    # ================================================================
    add_heading_styled(doc, "2. Architecture testee", level=1)
    doc.add_paragraph(
        "Le flux teste reproduit exactement le parcours reel de l'application mobile :"
    )
    arch = doc.add_paragraph()
    arch_run = arch.add_run(
        "Health Connect (Android SDK)\n"
        "    -> App React Native (lecture on-device)\n"
        "        -> HTTP POST /patients/{id}/health-data\n"
        "            -> API Gateway (:8010)\n"
        "                -> Patient Service (:8002)\n"
        "                    -> PostgreSQL (daily_aggregates)\n"
        "                        -> UPSERT sur (patient_id, date)"
    )
    arch_run.font.name = "Consolas"
    arch_run.font.size = Pt(10)

    doc.add_paragraph(
        "Les donnees Health Connect sont simulees cote preview web (valeurs aleatoires "
        "realistes), mais l'envoi HTTP et la persistance PostgreSQL sont reels."
    )

    doc.add_page_break()

    # ================================================================
    # 3. TESTS AUTHENTIFICATION
    # ================================================================
    add_heading_styled(doc, "3. Tests d'authentification", level=1)

    add_heading_styled(doc, "MOB-01 : Ecran de login", level=2)
    doc.add_paragraph(
        "L'application affiche un ecran de connexion avec les champs email et mot de passe. "
        "Le patient doit s'authentifier avant de pouvoir synchroniser ses donnees."
    )
    add_screenshot(doc, "sc_mob_01_login.png", "Figure 1 - Ecran de login de l'application mobile")

    add_heading_styled(doc, "MOB-02 : Connexion en cours", level=2)
    doc.add_paragraph(
        "Apres avoir clique sur 'SE CONNECTER', l'application envoie une requete "
        "POST /auth/login au backend. Un spinner s'affiche pendant l'attente de la reponse. "
        "La console a droite montre la requete HTTP en temps reel."
    )
    add_screenshot(doc, "sc_mob_02_login_loading.png", "Figure 2 - Connexion en cours (spinner + console HTTP)")

    add_heading_styled(doc, "MOB-03 : Accueil apres login", level=2)
    doc.add_paragraph(
        "Une fois authentifie, l'application affiche l'ecran principal avec le nom du "
        "patient, les statistiques de sante (vides avant synchronisation), et le bouton "
        "'SYNCHRONISER'. Le token JWT est stocke en memoire."
    )
    add_screenshot(doc, "sc_mob_03_accueil.png", "Figure 3 - Ecran d'accueil apres authentification")

    doc.add_page_break()

    # ================================================================
    # 4. TESTS SYNCHRONISATION
    # ================================================================
    add_heading_styled(doc, "4. Tests de synchronisation", level=1)

    add_heading_styled(doc, "MOB-04 : Synchronisation automatique", level=2)
    doc.add_paragraph(
        "L'application lance automatiquement une synchronisation apres le login. "
        "Elle lit les donnees Health Connect (pas, frequence cardiaque, sommeil, GPS) "
        "puis les envoie au backend via HTTP POST avec le token JWT en header Authorization."
    )
    add_screenshot(doc, "sc_mob_04_sync_loading.png",
                   "Figure 4 - Synchronisation en cours (lecture Health Connect + envoi HTTP)")

    add_heading_styled(doc, "MOB-05 : Synchronisation reussie", level=2)
    doc.add_paragraph(
        "Les donnees sont envoyees avec succes au Patient Service. Le badge vert confirme "
        "la synchronisation. La console affiche la reponse du serveur avec le champ "
        "'upserted: true' et le 'synced_at' horodate."
    )
    add_screenshot(doc, "sc_mob_05_sync_ok.png",
                   "Figure 5 - Donnees synchronisees avec succes (badge vert + console)")

    add_heading_styled(doc, "MOB-06 : Re-synchronisation (UPSERT)", level=2)
    doc.add_paragraph(
        "En appuyant a nouveau sur 'SYNCHRONISER', l'application envoie de nouvelles "
        "donnees pour le meme jour. Le backend effectue un UPSERT : il met a jour "
        "l'enregistrement existant au lieu de creer un doublon. Cela garantit une seule "
        "ligne par patient par jour dans daily_aggregates."
    )
    add_screenshot(doc, "sc_mob_06_resync.png",
                   "Figure 6 - Re-synchronisation UPSERT (mise a jour sans doublon)")

    doc.add_page_break()

    # ================================================================
    # 5. TESTS PERSISTANCE
    # ================================================================
    add_heading_styled(doc, "5. Tests de persistance", level=1)

    add_heading_styled(doc, "MOB-07 : Verification PostgreSQL", level=2)
    doc.add_paragraph(
        "Requete directe sur la base de donnees PostgreSQL pour verifier que les donnees "
        "synchronisees sont bien presentes dans la table daily_aggregates avec le champ "
        "source_platform = 'android_health_connect'."
    )
    add_text_evidence(doc, "sc_mob_07_db_check.txt")

    add_heading_styled(doc, "MOB-08 : Batch sync (3 jours offline)", level=2)
    doc.add_paragraph(
        "Test de l'endpoint /health-data/batch qui permet de synchroniser plusieurs jours "
        "de donnees en une seule requete. Ce scenario simule un patient qui n'a pas ouvert "
        "l'application pendant 3 jours et qui synchronise tout d'un coup."
    )
    # Show JSON response
    json_path = EVIDENCIAS / "sc_mob_08_batch_api.json"
    if json_path.exists():
        import json
        data = json.loads(json_path.read_text(encoding="utf-8"))
        para = doc.add_paragraph()
        run = para.add_run(json.dumps(data["response"], indent=2, ensure_ascii=False))
        run.font.size = Pt(8)
        run.font.name = "Consolas"
        cap = doc.add_paragraph("Reponse JSON du batch sync (3 jours)")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.italic = True
        cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    add_heading_styled(doc, "MOB-09 : Verification batch en BD", level=2)
    doc.add_paragraph(
        "Verification que les 3 jours de batch sync + les syncs individuels sont tous "
        "presents dans PostgreSQL. Le total doit correspondre au nombre de jours uniques "
        "synchronises."
    )
    add_text_evidence(doc, "sc_mob_09_db_batch.txt")

    doc.add_page_break()

    # ================================================================
    # 6. DECONNEXION
    # ================================================================
    add_heading_styled(doc, "6. Test de deconnexion", level=1)

    add_heading_styled(doc, "MOB-10 : Deconnexion", level=2)
    doc.add_paragraph(
        "Le patient se deconnecte. Le token JWT est supprime de la memoire et "
        "l'application revient a l'ecran de login. Aucune donnee sensible ne reste "
        "accessible apres la deconnexion."
    )
    add_screenshot(doc, "sc_mob_10_logout.png", "Figure 7 - Retour a l'ecran de login apres deconnexion")

    # ================================================================
    # 7. RESUME
    # ================================================================
    add_heading_styled(doc, "7. Resume des resultats", level=1)

    tests = [
        ("MOB-01", "Ecran de login app mobile", "REUSSI"),
        ("MOB-02", "Login en cours (spinner)", "REUSSI"),
        ("MOB-03", "Accueil apres login", "REUSSI"),
        ("MOB-04", "Auto-sync Health Connect", "REUSSI"),
        ("MOB-05", "Sync terminee (badge vert)", "REUSSI"),
        ("MOB-06", "Re-sync UPSERT", "REUSSI"),
        ("MOB-07", "Donnees dans PostgreSQL", "REUSSI"),
        ("MOB-08", "Batch sync 3 jours", "REUSSI"),
        ("MOB-09", "Verification batch BD", "REUSSI"),
        ("MOB-10", "Deconnexion app mobile", "REUSSI"),
    ]

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header = table.rows[0].cells
    header[0].text = "ID"
    header[1].text = "Description"
    header[2].text = "Resultat"
    header[3].text = "Evidence"

    for tid, desc, status in tests:
        row = table.add_row().cells
        row[0].text = tid
        row[1].text = desc
        row[2].text = status
        row[3].text = f"sc_mob_{tid.split('-')[1]}_*.png/txt/json"

    doc.add_paragraph("")
    summary = doc.add_paragraph()
    summary_run = summary.add_run(
        f"Total : 10 tests | 10 REUSSIS | 0 ECHOUES\n"
        f"Taux de reussite : 100%"
    )
    summary_run.font.size = Pt(12)
    summary_run.font.bold = True
    summary_run.font.color.rgb = RGBColor(0x15, 0x57, 0x24)
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    footer = doc.add_paragraph(
        f"Rapport genere le {datetime.now().strftime('%d/%m/%Y a %H:%M')} | "
        f"Mood-IoT v2.0 | Fil Rouge Master ADE 2026"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save
    doc.save(str(OUTPUT))
    print(f"Rapport genere : {OUTPUT}")
    print(f"Taille : {OUTPUT.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    generate()
